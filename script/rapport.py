import ast
import re
import numpy as np
import pandas as pd
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Optional, Dict

# --- Parse 'code' -> liste d'entiers (gère "[5, 1]" ou "5")
def _parse_code_list(val):
    if val is None or (isinstance(val, float) and np.isnan(val)):
        return []
    if isinstance(val, (list, tuple, set, np.ndarray)):
        try:
            return [int(x) for x in val]
        except Exception:
            return []
    s = str(val).strip()
    if not s:
        return []
    if s.startswith('[') and s.endswith(']'):
        try:
            parsed = ast.literal_eval(s)
            if isinstance(parsed, (list, tuple)):
                return [int(x) for x in parsed]
        except Exception:
            pass
    nums = re.findall(r"-?\d+", s)
    return [int(x) for x in nums] if nums else []

# --- Convertit en datetime : supporte Serial Excel ou texte
def _to_datetime_safely(s: pd.Series) -> pd.Series:
    if pd.api.types.is_datetime64_any_dtype(s):
        return s
    if pd.api.types.is_numeric_dtype(s):
        non_na = s.dropna()
        if len(non_na) and (non_na.gt(10_000).mean() > 0.9):
            return pd.to_datetime(s, unit="D", origin="1899-12-30", errors="coerce")
    converted = pd.to_datetime(s, errors="coerce")
    if converted.isna().all() and pd.api.types.is_numeric_dtype(s):
        converted = pd.to_datetime(s, unit="D", origin="1899-12-30", errors="coerce")
    return converted

# --- Définitions des codes (0..13)
CODE_DEFINITIONS = {
    0:  "Correcte",
    1:  "Moins de 10 minutes",
    2:  "Interactive > 6 heures",
    3:  "Interactive non active pour 30 minutes",
    4:  "Non active > 30 minutes",
    5:  "Inactive",
    6:  "A100 => MIG 1g.5bg",
    7:  "A100 => MIG 2g.10bg",
    8:  "A100 => MIG 3g.20bg",
    9:  "A100 trop de mémoire CPU",
    10: "MIG trop de mémoire CPU",
    11: "Tâche MIG qui demande > 1 MIG",
    12: "MIG 3g.20bg => MIG 2g.10bg",
    13: "MIG 2g.10bg => MIG 1g.5bg",
}

def build_jobs_report_doc(
    df: pd.DataFrame,
    *,
    date_col: str = "Start",     # utilisé pour déterminer le mois (YYYY-MM)
    state_col: str = "State",
    code_col: str = "code",
    output_path: str = "rapport_jobs.docx",
    code_definitions: Optional[Dict[int, str]] = None,
) -> str:
    """
    Génère un rapport Word (.docx) avec :
      - Global : total, % par code, % multi-codes, % par état
      - Par mois : mêmes stats pour chaque mois (selon 'date_col')

    Paramètres
    ----------
    df : DataFrame (doit ressembler à ton fichier d'exemple : colonnes 'Start', 'State', 'code', etc.)  # [1](https://ulavaldti-my.sharepoint.com/personal/hegin22_ulaval_ca/_layouts/15/Doc.aspx?sourcedoc=%7BAF372B23-C9A1-488A-A64E-F44E265C9D50%7D&file=job_frame-test-job_moins_10_min.xlsx&action=default&mobileredirect=true)
    date_col : colonne à utiliser pour le groupement mensuel (ex. 'Start' ou 'End')
    state_col : colonne d'état (ex. 'State')
    code_col : colonne des codes (ex. 'code')
    output_path : chemin du .docx à créer
    code_definitions : dict {code_int : description} (par défaut, dictionnaire 0..13 ci-dessus)

    Retour
    ------
    output_path : chemin du fichier .docx généré
    """
    if code_definitions is None:
        code_definitions = CODE_DEFINITIONS

    data = pd.DataFrame(df).copy()
    # Vérifications minimales
    for col in [date_col, state_col, code_col]:
        if col not in data.columns:
            raise ValueError(f"Colonne manquante: '{col}'")

    # Codes -> liste d'entiers
    data[code_col] = data[code_col].apply(_parse_code_list)

    # Date -> mois (YYYY-MM)
    data[date_col] = _to_datetime_safely(data[date_col])
    data["__month__"] = data[date_col].dt.to_period("M").astype(str)

    # Normaliser State
    data[state_col] = data[state_col].astype(str).fillna("")

    total_all = len(data)
    all_codes = list(sorted(code_definitions.keys()))

    def stats_for_subset(sub: pd.DataFrame):
        total = len(sub)
        # % par code : un job avec codes [1,5] compte pour 1 dans 1 et 1 dans 5
        code_counts = {
            c: (int(sub[code_col].apply(lambda L: c in L).sum()),
                (sub[code_col].apply(lambda L: c in L).mean() * 100.0) if total else 0.0)
            for c in all_codes
        }
        # % multi-codes
        multi_mask = sub[code_col].apply(lambda L: len(L) > 1)
        multi_cnt = int(multi_mask.sum())
        multi_pct = (multi_mask.mean() * 100.0) if total else 0.0
        # % par état
        vc = sub[state_col].value_counts(dropna=False)
        state_stats = {k: (int(n), (n / total * 100.0) if total else 0.0) for k, n in vc.items()}
        return {"total": total, "code_stats": code_counts, "multi": (multi_cnt, multi_pct), "state_stats": state_stats}

    # Global
    global_stats = stats_for_subset(data)
    # Par mois
    months = sorted(m for m in data["__month__"].dropna().unique())
    by_month = {m: stats_for_subset(data[data["__month__"] == m]) for m in months}

    # --- Construction du document
    doc = Document()

    # Titre
    doc.add_heading("Rapport des tâches — Synthèse par mois et globale", level=0)
    p = doc.add_paragraph(f"Généré le {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Légende des codes (+ % global par code)
    doc.add_heading("Légende des codes", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    hdr[0].text = "Code"
    hdr[1].text = "Définition"
    hdr[2].text = "Présent (%) — Global"
    for c in all_codes:
        row = table.add_row().cells
        row[0].text = str(c)
        row[1].text = code_definitions.get(c, "Inconnu")
        cnt, pct = global_stats["code_stats"][c]
        row[2].text = f"{pct:,.1f}% ({cnt}/{total_all})"

    doc.add_paragraph("")

    # Global
    doc.add_heading("Statistiques globales", level=1)
    doc.add_paragraph(f"Nombre total de tâches : {global_stats['total']}")

    doc.add_heading("Répartition par code (global)", level=2)
    t_code = doc.add_table(rows=1, cols=4)
    hc = t_code.rows[0].cells
    hc[0].text = "Code"; hc[1].text = "Définition"; hc[2].text = "Tâches (#)"; hc[3].text = "Tâches (%)"
    for c in all_codes:
        cnt, pct = global_stats["code_stats"][c]
        r = t_code.add_row().cells
        r[0].text = str(c); r[1].text = code_definitions.get(c, "Inconnu"); r[2].text = f"{cnt}"; r[3].text = f"{pct:,.1f}%"

    mc_cnt, mc_pct = global_stats["multi"]
    doc.add_paragraph("")
    doc.add_heading("Tâches multi-codes (>1 code) — global", level=2)
    t_mc = doc.add_table(rows=2, cols=2)
    t_mc.rows[0].cells[0].text = "Tâches (#)"
    t_mc.rows[0].cells[1].text = "Tâches (%)"
    t_mc.rows[1].cells[0].text = str(mc_cnt)
    t_mc.rows[1].cells[1].text = f"{mc_pct:,.1f}%"

    doc.add_paragraph("")
    doc.add_heading("Répartition par état (global)", level=2)
    t_state = doc.add_table(rows=1, cols=3)
    hs = t_state.rows[0].cells
    hs[0].text = "État"; hs[1].text = "Tâches (#)"; hs[2].text = "Tâches (%)"
    for st, (cnt, pct) in sorted(global_stats["state_stats"].items(), key=lambda kv: (-kv[1][0], kv[0])):
        row = t_state.add_row().cells
        row[0].text = str(st); row[1].text = str(cnt); row[2].text = f"{pct:,.1f}%"

    # Par mois (une page par mois)
    for m in months:
        doc.add_page_break()
        doc.add_heading(f"Statistiques — {m}", level=1)
        st = by_month[m]
        doc.add_paragraph(f"Nombre total de tâches : {st['total']}")

        doc.add_heading("Répartition par code", level=2)
        t_code_m = doc.add_table(rows=1, cols=4)
        hcm = t_code_m.rows[0].cells
        hcm[0].text = "Code"; hcm[1].text = "Définition"; hcm[2].text = "Tâches (#)"; hcm[3].text = "Tâches (%)"
        for c in all_codes:
            cnt, pct = st["code_stats"][c]
            row = t_code_m.add_row().cells
            row[0].text = str(c); row[1].text = code_definitions.get(c, "Inconnu"); row[2].text = f"{cnt}"; row[3].text = f"{pct:,.1f}%"

        doc.add_paragraph("")
        doc.add_heading("Tâches multi-codes (>1 code)", level=2)
        t_mc_m = doc.add_table(rows=2, cols=2)
        t_mc_m.rows[0].cells[0].text = "Tâches (#)"; t_mc_m.rows[0].cells[1].text = "Tâches (%)"
        t_mc_m.rows[1].cells[0].text = str(st["multi"][0]); t_mc_m.rows[1].cells[1].text = f"{st['multi'][1]:,.1f}%"

        doc.add_paragraph("")
        doc.add_heading("Répartition par état", level=2)
        t_state_m = doc.add_table(rows=1, cols=3)
        hsm = t_state_m.rows[0].cells
        hsm[0].text = "État"; hsm[1].text = "Tâches (#)"; hsm[2].text = "Tâches (%)"
        for st_name, (cnt, pct) in sorted(st["state_stats"].items(), key=lambda kv: (-kv[1][0], kv[0])):
            row = t_state_m.add_row().cells
            row[0].text = str(st_name); row[1].text = str(cnt); row[2].text = f"{pct:,.1f}%"

    doc.save(output_path)
    return output_path